import pandas as pd
import docx2txt
import os
from docx import Document
from docx.shared import Pt
import glob
import csv
import shutil
import re
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
import zipfile
import tempfile


# convert the docx file to txt
def convert_doc_txt(docx_file, txt_file):
    MY_TEXT = docx2txt.process(docx_file)
    with open(txt_file, "w") as text_file:
        print(MY_TEXT, file=text_file)


# Step 2
# Convert xlsx file to csv
def xlsx_to_csv(xlsx_file, csv_file):
    df = pd.read_excel(xlsx_file)
    # pd.read_csv(csv_file, skiprows=3)
    df.to_csv(csv_file, index=False)



#  clean the csv file
def clean_csv_file(input_csv, output_csv):
    df = pd.read_csv(input_csv, skiprows=2) #names=["Donor Name", "Donor Amount", "Add Amount"])
    columns_to_keep = ['Donor Name', 'Donor Amount', 'Add Amount']
    df = df[columns_to_keep]
    df.to_csv(output_csv, index=False)


# read csv file for computation
def read_csv(filename):
    names = []
    amounts = []
    # add_amounts = []

    with open(filename, mode='r', newline='', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        
        # Check if the headers 'Donor Name' and 'Donor Amount' exist in the CSV file
        if 'Donor Name' in reader.fieldnames and 'Donor Amount' in reader.fieldnames:
            for row in reader:
                name = row['Donor Name'].strip()
                amount_str = row['Donor Amount'].strip()
                # add_amount = row['Add Amount'].strip()
                # print(add_amount)
                # Handle missing or empty amount values
                if amount_str:
                    amounts.append(amount_str)
                else:
                    amounts.append("0")
                
                # add_amounts.append(add_amount)
                names.append(name)
        else:
            print("CSV file does not contain the expected column headers.")

    return names, amounts

def replace_words(name):
    # Read the text file
    with open('sample.txt', 'r') as file:
        content = file.read()

    # Replace the "donor name" and "donor amount" with your desired text
    content_modified = content.replace('<Friends>', name)
    # content_modified = content_modified.replace('Donor Amount', amount)

    # Create a new file with a dynamic name (e.g., name.txt) and store the modified content there
    new_filename = f'{name}.txt'
    with open(new_filename, 'w') as new_file:
        new_file.write(content_modified)

    print(f'Text successfully replaced and saved to {new_filename}!')


def replace_content_after_date(docx_filename, txt_filename, output_filename, add):
    # Read the content from the text file
    with open(txt_filename, 'r', encoding='utf-8') as txt_file:
        new_content = txt_file.read()

        
    # Open the existing .docx file
    doc = Document(docx_filename)
    
    # print(add)
    if add == True:
        doc.save(output_filename)
    # Flag to indicate when to start replacing content
    start_replacing = False
    
    # Iterate through paragraphs and replace content after finding the date
    for paragraph in doc.paragraphs:
        if "27th September 2023" in paragraph.text:
            start_replacing = True
        
        # If the flag is set, clear the paragraph and add the new content
        if start_replacing:
            for run in paragraph.runs:
                run.clear()
            paragraph.add_run(new_content)
        break
    # Save the modified document with a new name
    doc.save(output_filename)

    print(f'Content after "27th September 2023" has been replaced in {output_filename}.')


def replace_words_in_docx(docx_filename, replacements):
    # Load the .docx document
    doc = Document(docx_filename)

    # Iterate through paragraphs
    for paragraph in doc.paragraphs:
        for old_word, new_word in replacements.items():
            # Replace the old word with the new word in the paragraph
            paragraph.text = paragraph.text.replace(old_word, new_word)

    # Save the modified document with a new name
    output_filename = f"{docx_filename}"
    doc.save(output_filename)

    print(f"Words replaced and saved in {output_filename}")



def move_docx_files_to_thank_you_folder(source_folder, destination_folder):
    # Ensure the "Thank you" folder exists; create it if it doesn't
    if not os.path.exists(destination_folder):
        os.makedirs(destination_folder)

    # Iterate through files in the source folder
    for filename in os.listdir(source_folder):
        source_path = os.path.join(source_folder, filename)

        # Check if the file is a .docx file
        if filename.endswith(".docx"):
            destination_path = os.path.join(destination_folder, filename)
            if filename == "no_amount_template.docx":
                continue
            if filename == "amount_template.docx":
                continue


            # Move the .docx file to the "Thank you" folder
            shutil.move(source_path, destination_path)
            print(f'Moved {filename} to {destination_folder}')
    




def store_add_amount_users(csvfile):
   with open(csvfile, mode='r') as csv_file:
    csv_reader = csv.DictReader(csv_file)
    
    # Create a list to store rows that match the condition
    filtered_rows = []
    
    # Loop through each row in the original CSV
    for row in csv_reader:
        if row['Add Amount'] == 'True':
            # Append the row to the list
            filtered_rows.append(row)

# Create a new CSV file for storing names with "Add Amount" equal to "True"
    with open('filtered_names.csv', mode='w', newline='') as filtered_csv_file:
        fieldnames = ['Donor Name', 'Donor Amount', 'Add Amount']
        writer = csv.DictWriter(filtered_csv_file, fieldnames=fieldnames)
        writer.writeheader()
        
        # Write the filtered rows to the new CSV file
        writer.writerows(filtered_rows)



def extract_all_donor_name(csvfile):
    names = []
    with open(csvfile, mode='r' ,newline='', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        if "Donor Name" in reader.fieldnames:
            for row in reader:
                name = row['Donor Name'].strip()
                names.append(name)
        else:
            print("CSV file does not contain the expected column headers.")

    return names



def delete_all_txt_files(directory_path):
    try:
        # Ensure the directory path exists
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        # List all TXT files in the directory
        txt_files = [file for file in os.listdir(directory_path) if file.endswith(".txt")]

        for txt_file in txt_files:
            # Construct the full file path
            file_path = os.path.join(directory_path, txt_file)
            if txt_file == "requirements.txt":
                continue
            # Delete the file
            os.remove(file_path)

            print(f"Deleted: {txt_file}")

        print(f"Deleted {len(txt_files)} TXT files.")
    except Exception as e:
        print(f"An error occurred: {e}")




# add api endpoints and handler
app = FastAPI()

origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
@app.post("/upload/")
async def upload_files(no_docx_file: UploadFile = File(...), xlsx_file: UploadFile = File(...)):
    try:
        # Create a directory to store uploaded files (if it doesn't exist)
        # if not os.path.exists("uploads"):
        #     os.makedirs("uploads")

        # Save the uploaded DOCX file
        docx_file_path = os.path.join("./", "no_amount_template.docx")
        with open(docx_file_path, "wb") as f:
            f.write(no_docx_file.file.read())
        # docx_file_path = os.path.join("./", "no_amount_template.docx")
        # with open(docx_file_path, "wb") as f:
        #     f.write(no_docx_file.file.read())

        # Save the uploaded XLSX file
        xlsx_file_path = os.path.join("./", "donations.xlsx")
        with open(xlsx_file_path, "wb") as f:
            f.write(xlsx_file.file.read())
        # _,_, add_amount = xlsx_to_csv("donations.xlsx", "donations.csv")

        xlsx_to_csv("donations.xlsx", "donations.csv")
        convert_doc_txt("no_amount_template.docx", "sample.txt")
        clean_csv_file("donations.csv", "donations.csv")
        store_add_amount_users("donations.csv")

        # names, amounts = read_csv("donations.csv")
        donor_names = extract_all_donor_name("filtered_names.csv")
        for name in donor_names:
                text_file = name+".txt"
                # docx_folder = "completed/"+ name + "." + "docx"
                docx_file = name + ".docx"
                # cleaned_amount = amount.replace('₦', '').replace(',', '')
        
                # try:
                #     amount = int(cleaned_amount)
                # except ValueError:
                #     # Handle invalid amount value here
                #     print(f"Invalid amount value for {name}: {amount}")
                #     continue

                # Format the amount with commas
                # formatted_amount = '{:,.0f}'.format(amount)

                # formatted_amount = '{:,}'.format(int(amount))
                replacements = {
                    "<Friends>": name}
                replace_words(name)
                replace_content_after_date("no_amount_template.docx", text_file, docx_file, "")
                replace_words_in_docx(docx_file, replacements)
    #             # replace_by_add_amount_flag(filename)
        
        # open_docx_file("./completed")
        # donor_names = extract_all_donor_name("filtered_names.csv")
        # for name in donor_names:
        #     # # root = os.path.abspath("completed")
        #     # # file_path = f'{root}/{name}.docx'
        #     # file = f'{name}.docx'
        #     # file_path = os.path.join(".", file)
        #     # if os.path.exists(file_path):
        #     #     os.remove(file_path)
        #     # else:
        #     #     print("file not found")
        #     text_file = f'{name}.txt'
        #     # d_file = f'{name}.docx'
        #     d_file = name + ".docx"
        #     replacements = {
        #         "<Friends>": name
        #     }
        #     # replace_words(name)
        #     # replace_content_after_date("no_amount_template.docx", text_file, docx_file, "")
        #     # replace_words_in_docx(docx_file, replacements)
        #     if os.path.exists(d_file):
        #         replace_words(name)
        #         replace_content_after_date("no_amount_template.docx", text_file, d_file, "")
        #         replace_words_in_docx(d_file, replacements)
        #     else:
        #         print(f"Docx file not found for {name}")

        delete_all_txt_files(".")
        move_docx_files_to_thank_you_folder(".", "no_amount_completed")
        return JSONResponse(content={
            "message": "Files uploaded successfully and modified",
            "download_link": "https://doc-no-amount-server.onrender.com/download/no_amount_completed"
            }, status_code=200)
    except Exception as e:
         return JSONResponse(content={"error": str(e)}, status_code=500)

@app.get("/download/{folder_name}")
async def download_folder(folder_name: str):
    try:
        folder_path = os.path.join("./", folder_name)

        if not os.path.exists(folder_path):
            raise HTTPException(status_code=404, detail="Folder not found")

        # Create a temporary directory to store the zip archive
        temp_dir = tempfile.mkdtemp()

        # Create a zip file to store the folder contents
        zip_filename = f"{folder_name}.zip"
        zip_filepath = os.path.join(temp_dir, zip_filename)

        with zipfile.ZipFile(zip_filepath, "w", zipfile.ZIP_DEFLATED) as zipf:
            for root, _, files in os.walk(folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, folder_path)
                    zipf.write(file_path, arcname=arcname)

        # Serve the zip archive for download
        return FileResponse(zip_filepath, media_type='application/zip', filename=zip_filename)
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8002)
